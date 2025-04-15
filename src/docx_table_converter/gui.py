"""
GUI application for the DOCX Table Converter.
"""

import sys
import os
import pandas as pd
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                            QHBoxLayout, QPushButton, QLabel, QFileDialog,
                            QSpinBox, QComboBox, QTextEdit, QTableWidget,
                            QTableWidgetItem, QCheckBox, QGroupBox, QMessageBox,
                            QDialog, QLineEdit, QFormLayout, QHeaderView, QSplitter,
                            QTabWidget, QGridLayout, QDoubleSpinBox)
from PyQt5.QtCore import Qt, QSize, QTimer
from PyQt5.QtGui import QFont, QIcon
from docx_table_converter.core import (write_table_to_docx, read_table_from_file, parse_clipboard_data,
                  write_tables_to_docx)
from docx import Document
import tempfile
import io
import shutil

# 翻译字典
TRANSLATIONS = {
    "window_title": {
        "zh": "论文三线表--一键转换",
        "en": "Research Table Converter"
    },
    "single_table": {
        "zh": "单表转换",
        "en": "Single Table"
    },
    "batch_tables": {
        "zh": "批量转换",
        "en": "Batch Tables"
    },
    "select_file": {
        "zh": "选择文件",
        "en": "Select File"
    },
    "paste_data": {
        "zh": "粘贴数据",
        "en": "Paste Data"
    },
    "preview": {
        "zh": "预览",
        "en": "Preview"
    },
    "export": {
        "zh": "导出",
        "en": "Export"
    },
    "header_rows": {
        "zh": "表头行数",
        "en": "Header Rows"
    },
    "index_column": {
        "zh": "索引列",
        "en": "Index Column"
    },
    "include_index": {
        "zh": "包含索引",
        "en": "Include Index"
    },
    "font": {
        "zh": "字体",
        "en": "Font"
    },
    "font_size": {
        "zh": "字号",
        "en": "Font Size"
    },
    "table_title": {
        "zh": "表格标题",
        "en": "Table Title"
    },
    "table_description": {
        "zh": "表格描述",
        "en": "Table Description"
    },
    "none": {
        "zh": "无",
        "en": "None"
    },
    "import_success": {
        "zh": "导入成功",
        "en": "Import Success"
    },
    "paste_success": {
        "zh": "粘贴成功",
        "en": "Paste Success"
    },
    "no_data": {
        "zh": "请先加载或粘贴有效的表格数据！",
        "en": "Please load or paste valid table data first!"
    },
    "save_docx": {
        "zh": "保存DOCX文件",
        "en": "Save DOCX File"
    },
    "export_success": {
        "zh": "导出成功",
        "en": "Export Success"
    },
    "select_files": {
        "zh": "选择文件",
        "en": "Select Files"
    },
    "pasted_data": {
        "zh": "粘贴数据",
        "en": "Pasted Data"
    },
    "table": {
        "zh": "表格",
        "en": "Table"
    },
    "delete": {
        "zh": "删除",
        "en": "Delete"
    },
    "no_tables": {
        "zh": "请先添加表格！",
        "en": "Please add tables first!"
    },
    "merge_output": {
        "zh": "合并输出",
        "en": "Merge Output"
    },
    "select_output_dir": {
        "zh": "选择输出目录",
        "en": "Select Output Directory"
    },
    "add_files": {
        "zh": "添加文件",
        "en": "Add Files"
    },
    "paste_tables": {
        "zh": "粘贴表格",
        "en": "Paste Tables"
    },
    "clear_all": {
        "zh": "清空全部",
        "en": "Clear All"
    },
    "preview_all": {
        "zh": "预览全部",
        "en": "Preview All"
    },
    "file_path": {
        "zh": "文件路径/数据来源",
        "en": "File Path/Data Source"
    },
    "operations": {
        "zh": "操作",
        "en": "Operations"
    },
    "clipboard_empty": {
        "zh": "剪贴板为空！",
        "en": "Clipboard is empty!"
    },
    "parse_error": {
        "zh": "解析数据时出错：{0}",
        "en": "Parse error: {0}"
    },
    "cancel": {
        "zh": "取消",
        "en": "Cancel"
    },
    "close": {
        "zh": "关闭",
        "en": "Close"
    },
    "save_success": {
        "zh": "保存成功",
        "en": "Save Success"
    },
    "info": {
        "zh": "提示",
        "en": "Information"
    },
    "warning": {
        "zh": "警告",
        "en": "Warning"
    },
    "error": {
        "zh": "错误",
        "en": "Error"
    },
    "clear": {
        "zh": "清除",
        "en": "Clear"
    },
    "load_error": {
        "zh": "加载文件数据时出错：{0}",
        "en": "Error loading file data: {0}"
    }
}

def get_translation(key, language="zh"):
    """获取指定语言的翻译"""
    if key in TRANSLATIONS:
        return TRANSLATIONS[key].get(language, key)
    return key

def show_message(msg_type, text):
    """显示消息对话框
    
    Args:
        msg_type: 消息类型，可选值：info, warning, error
        text: 消息内容
    """
    if msg_type == "info":
        QMessageBox.information(None, get_translation("info"), text)
    elif msg_type == "warning":
        QMessageBox.warning(None, get_translation("warning"), text)
    elif msg_type == "error":
        QMessageBox.critical(None, get_translation("error"), text)

def create_button(text, callback=None):
    """创建按钮
    
    Args:
        text: 按钮文本
        callback: 点击回调函数
    
    Returns:
        QPushButton: 按钮对象
    """
    btn = QPushButton(text)
    if callback:
        btn.clicked.connect(callback)
    return btn

def create_combo_box(items=None, default_index=0):
    """创建下拉框
    
    Args:
        items: 选项列表
        default_index: 默认选中项索引
    
    Returns:
        QComboBox: 下拉框对象
    """
    combo = QComboBox()
    if items:
        combo.addItems(items)
        combo.setCurrentIndex(default_index)
    return combo

def create_spin_box(min_value, max_value, default_value):
    """创建数字输入框
    
    Args:
        min_value: 最小值
        max_value: 最大值
        default_value: 默认值
    
    Returns:
        QSpinBox: 数字输入框对象
    """
    spin = QSpinBox()
    spin.setMinimum(min_value)
    spin.setMaximum(max_value)
    spin.setValue(default_value)
    return spin

def create_double_spin_box(min_value, max_value, default_value):
    """创建双精度浮点数输入框
    
    Args:
        min_value: 最小值
        max_value: 最大值
        default_value: 默认值
    
    Returns:
        QDoubleSpinBox: 双精度浮点数输入框对象
    """
    spin = QDoubleSpinBox()
    spin.setMinimum(min_value)
    spin.setMaximum(max_value)
    spin.setValue(default_value)
    return spin

def setup_table_widget(headers, col_widths=None, stretch_cols=None):
    """设置表格部件
    
    Args:
        headers: 表头列表
        col_widths: 列宽列表
        stretch_cols: 自动拉伸的列索引列表
    
    Returns:
        QTableWidget: 表格部件对象
    """
    table = QTableWidget()
    table.setColumnCount(len(headers))
    table.setHorizontalHeaderLabels(headers)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            table.setColumnWidth(i, width)
    
    if stretch_cols:
        header = table.horizontalHeader()
        for col in stretch_cols:
            header.setSectionResizeMode(col, QHeaderView.Stretch)
    
    return table

def load_dataframe(file_path, sheet_name=None, header_rows=1):
    """加载DataFrame
    
    Args:
        file_path: 文件路径
        sheet_name: Excel工作表名称
        header_rows: 表头行数
    
    Returns:
        pd.DataFrame: 数据表格
    """
    try:
        return read_table_from_file(
            file_path,
            sheet_name=sheet_name,
            header_rows=header_rows
        )
    except Exception as e:
        raise Exception(get_translation("load_error").format(str(e)))

def create_temp_docx():
    """
    Creates a temporary DOCX file and ensures it exists.
    
    Returns:
        str: Path to the temporary DOCX file
    """
    try:
        # Create a temporary file with .docx extension
        temp_dir = tempfile.gettempdir()
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', dir=temp_dir, delete=False)
        temp_path = temp_file.name
        temp_file.close()
        
        # Create an empty document to ensure the file exists
        doc = Document()
        doc.save(temp_path)
        
        if not os.path.exists(temp_path):
            raise FileNotFoundError(f"Failed to create temporary file at {temp_path}")
            
        return temp_path
    except Exception as e:
        QMessageBox.critical(None, "错误", f"创建临时文件失败: {str(e)}")
        return None

def cleanup_temp_file(file_path):
    """
    Safely deletes a temporary file.
    
    Args:
        file_path (str): Path to the temporary file
    """
    try:
        if file_path and os.path.exists(file_path):
            os.unlink(file_path)
    except Exception as e:
        print(f"Warning: Failed to delete temporary file {file_path}: {str(e)}")

class PasteTableDialog(QDialog):
    """粘贴表格数据对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.df = None
        self.header_rows = 1
        self.setup_ui()
    
    def setup_ui(self):
        """设置UI"""
        self.setWindowTitle(get_translation("paste_data"))
        self.setMinimumSize(1200, 800)
        
        layout = QVBoxLayout(self)
        
        # 表格部件
        self.table = QTableWidget()
        self.table.setColumnCount(10)  # 默认10列
        self.table.setRowCount(100)  # 默认100行
        layout.addWidget(self.table)
        
        # 表头行数
        header_layout = QHBoxLayout()
        header_label = QLabel(get_translation("header_rows"))
        self.header_spin = create_spin_box(0, 2, 1)
        header_layout.addWidget(header_label)
        header_layout.addWidget(self.header_spin)
        header_layout.addStretch()
        layout.addLayout(header_layout)
        
        # 按钮
        button_layout = QHBoxLayout()
        
        paste_btn = create_button(get_translation("paste"), self.paste_data)
        button_layout.addWidget(paste_btn)
        
        clear_btn = create_button(get_translation("clear"), self.clear_data)
        button_layout.addWidget(clear_btn)
        
        button_layout.addStretch()
        
        ok_btn = create_button("OK", self.accept)
        button_layout.addWidget(ok_btn)
        
        cancel_btn = create_button(get_translation("cancel"), self.reject)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
    
    def paste_data(self):
        """粘贴数据"""
        try:
            clipboard = QApplication.clipboard()
            text = clipboard.text()
            
            if not text.strip():
                show_message("warning", get_translation("clipboard_empty"))
                return
            
            # 尝试解析数据
            rows = text.strip().split('\n')
            data = [row.split('\t') for row in rows]
            
            # 获取最大列数
            max_cols = max(len(row) for row in data)
            
            # 设置表格大小
            self.table.setColumnCount(max_cols)
            self.table.setRowCount(len(data))
            
            # 填充数据
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    item = QTableWidgetItem(cell)
                    self.table.setItem(i, j, item)
            
            # 调整列宽
            self.table.resizeColumnsToContents()
            
            show_message("info", get_translation("paste_success"))
        except Exception as e:
            show_message("error", get_translation("parse_error").format(str(e)))
    
    def clear_data(self):
        """清空数据"""
        self.table.clearContents()
        self.table.setColumnCount(10)
        self.table.setRowCount(100)
    
    def get_dataframe(self):
        """获取DataFrame"""
        try:
            # 获取非空单元格的行列范围
            rows = self.table.rowCount()
            cols = self.table.columnCount()
            
            # 找到最后一个非空单元格的位置
            last_row = -1
            last_col = -1
            
            for row in range(rows):
                for col in range(cols):
                    item = self.table.item(row, col)
                    if item and item.text().strip():
                        last_row = max(last_row, row)
                        last_col = max(last_col, col)
            
            if last_row == -1 or last_col == -1:
                return None
            
            # 创建数据列表
            data = []
            for row in range(last_row + 1):
                row_data = []
                for col in range(last_col + 1):
                    item = self.table.item(row, col)
                    row_data.append(item.text() if item else '')
                data.append(row_data)
            
            # 转换为DataFrame
            self.header_rows = self.header_spin.value()
            if self.header_rows > 0:
                headers = data[:self.header_rows]
                if len(headers) == 1:
                    df = pd.DataFrame(data[self.header_rows:], columns=headers[0])
                else:
                    # 处理多级表头
                    df = pd.DataFrame(data[self.header_rows:])
                    for i, header_row in enumerate(headers):
                        df.columns = pd.MultiIndex.from_arrays([header_row] + [df.columns] * (i + 1))
            else:
                df = pd.DataFrame(data)
            
            return df
        except Exception as e:
            show_message("error", get_translation("parse_error").format(str(e)))
            return None


class DataSourceDialog(QDialog):
    def __init__(self, file_path, sheet_name, is_excel, header_rows, parent=None):
        super().__init__(parent)
        self.file_path = file_path
        self.sheet_name = sheet_name
        self.is_excel = is_excel
        self.header_rows = header_rows
        self.df = None
        self.raw_df = None
        self.setWindowTitle("数据源选择")
        self.setMinimumSize(800, 600)
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # 说明标签
        instruction_label = QLabel("您同时选择了文件和粘贴了数据，请选择使用哪一个数据源:")
        layout.addWidget(instruction_label)
        
        # 状态标签（用于显示临时提示）
        self.status_label = QLabel("")
        self.status_label.setAlignment(Qt.AlignCenter)
        self.status_label.setStyleSheet("color: green;")
        layout.addWidget(self.status_label)
        
        # 按钮布局
        button_layout = QHBoxLayout()
        
        # 文件数据按钮
        file_btn = QPushButton("使用文件数据")
        file_btn.clicked.connect(self.load_file_data)
        button_layout.addWidget(file_btn)
        
        # 粘贴数据按钮
        paste_btn = QPushButton("使用粘贴数据")
        paste_btn.clicked.connect(self.use_pasted_data)
        button_layout.addWidget(paste_btn)
        
        # 取消按钮
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
        
        # 表格控件
        self.table = QTableWidget()
        self.table.setRowCount(10)  # 默认行数
        self.table.setColumnCount(5)  # 默认列数
        
        # 设置表头
        headers = [f"列 {i+1}" for i in range(5)]
        self.table.setHorizontalHeaderLabels(headers)
        
        # 调整列宽
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.table)
        
        # 表头行数选择
        header_layout = QHBoxLayout()
        header_layout.addWidget(QLabel("表头行数:"))
        self.header_rows_spin = QSpinBox()
        self.header_rows_spin.setMinimum(0)
        self.header_rows_spin.setMaximum(2)
        self.header_rows_spin.setValue(self.header_rows)
        self.header_rows_spin.valueChanged.connect(self.update_table_headers)
        header_layout.addWidget(self.header_rows_spin)
        header_layout.addStretch()
        layout.addLayout(header_layout)
        
        # 索引列选择
        index_layout = QHBoxLayout()
        index_layout.addWidget(QLabel("索引列:"))
        self.index_column = QComboBox()
        self.index_column.addItem("无")
        self.index_column.currentIndexChanged.connect(self.update_dataframe)
        index_layout.addWidget(self.index_column)
        index_layout.addStretch()
        layout.addLayout(index_layout)
        
        self.setLayout(layout)

    def load_file_data(self):
        """加载文件数据"""
        try:
            # 直接使用 core 函数读取，传入正确的 header_rows
            self.df = read_table_from_file(
                self.file_path,
                sheet_name=self.sheet_name,
                header_rows=self.header_rows
            )
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载文件数据时出错: {str(e)}")
            self.reject()
    
    def use_pasted_data(self):
        """使用粘贴的数据"""
        self.reject()  # 拒绝对话框，表示不使用文件数据
    
    def get_dataframe(self):
        """获取当前DataFrame"""
        return self.df


class BatchProcessDialog(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.tables = []  # 存储粘贴的表格数据
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout()
        layout.setSpacing(10)  # 设置布局间距
        
        # ===== 1. 表格列表区域 =====
        tables_group = QGroupBox("表格列表")
        tables_layout = QVBoxLayout()
        tables_layout.setSpacing(10)
        
        # 表格列表控件
        self.table_list = QTableWidget()
        self.table_list.setColumnCount(4)  # 减少列数，合并标题和描述
        self.table_list.setHorizontalHeaderLabels([
            "文件路径",
            "表格标题",
            "表头行数",
            "操作"
        ])
        
        # 设置列宽策略
        header = self.table_list.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)  # 文件路径列自适应
        header.setSectionResizeMode(1, QHeaderView.Stretch)  # 标题列自适应
        header.setSectionResizeMode(2, QHeaderView.Fixed)   # 表头行数列固定
        header.setSectionResizeMode(3, QHeaderView.Fixed)   # 操作列固定
        
        # 设置默认列宽
        self.table_list.setColumnWidth(2, 120)   # 表头行数列宽度
        self.table_list.setColumnWidth(3, 120)  # 操作列宽度
        
        tables_layout.addWidget(self.table_list)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        add_file_btn = QPushButton("添加文件")
        add_file_btn.clicked.connect(self.add_files)
        button_layout.addWidget(add_file_btn)
        
        paste_data_btn = QPushButton("粘贴表格")
        paste_data_btn.clicked.connect(self.paste_table)
        button_layout.addWidget(paste_data_btn)
        
        remove_all_btn = QPushButton("清空列表")
        remove_all_btn.clicked.connect(self.clear_tables)
        button_layout.addWidget(remove_all_btn)
        
        tables_layout.addLayout(button_layout)
        tables_group.setLayout(tables_layout)
        layout.addWidget(tables_group)
        
        # ===== 2. 格式设置区域 =====
        format_group = QGroupBox("格式设置")
        format_layout = QGridLayout()
        format_layout.setHorizontalSpacing(100)  # 减小水平间距
        format_layout.setVerticalSpacing(10)    # 减小垂直间距
        
        # 中文字体
        format_layout.addWidget(QLabel("中文字体"), 0, 0)
        self.chinese_font_combo = create_combo_box(["宋体", "黑体", "楷体", "微软雅黑"])
        format_layout.addWidget(self.chinese_font_combo, 0, 1)
        
        # 英文字体
        english_label = QLabel("英文字体")
        format_layout.addWidget(english_label, 0, 2)
        self.english_font_combo = create_combo_box(["Times New Roman", "Arial", "Calibri"])
        format_layout.addWidget(self.english_font_combo, 0, 3)        
        # 字号
        format_layout.addWidget(QLabel("字号"), 1, 0)
        self.font_size_spin = create_spin_box(8, 72, 12)
        format_layout.addWidget(self.font_size_spin, 1, 1)
        
        # 三线表线宽
        format_layout.addWidget(QLabel("三线表线宽"), 1, 2)
        self.border_width_spin = QDoubleSpinBox()
        self.border_width_spin.setRange(0.5, 3.0)
        self.border_width_spin.setSingleStep(0.25)
        self.border_width_spin.setValue(1.0)
        format_layout.addWidget(self.border_width_spin, 1, 3)
        
        format_group.setLayout(format_layout)
        layout.addWidget(format_group)
        
        # ===== 3. 输出设置区域 =====
        output_group = QGroupBox("输出设置")
        output_layout = QGridLayout()
        output_layout.setHorizontalSpacing(100)  # 减小水平间距
        output_layout.setVerticalSpacing(10)    # 减小垂直间距
        
        # 标题加粗选项
        self.bold_caption_check = QCheckBox("标题加粗")
        self.bold_caption_check.setChecked(True)
        output_layout.addWidget(self.bold_caption_check, 0, 0)
        
        # 输出模式
        output_layout.addWidget(QLabel("输出模式"), 0, 1)
        self.output_mode = QComboBox()
        self.output_mode.addItems(["合并为一个文件", "每个表格单独保存"])
        output_layout.addWidget(self.output_mode, 0, 2)
        output_group.setLayout(output_layout)
        layout.addWidget(output_group)
        
        # ===== 4. 预览和导出按钮 =====
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        preview_btn = QPushButton("预览")
        preview_btn.clicked.connect(self.preview_tables)
        button_layout.addWidget(preview_btn)
        
        export_btn = QPushButton("导出")
        export_btn.clicked.connect(self.export_tables)
        button_layout.addWidget(export_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)

    def add_files(self):
        """添加文件"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择文件",
            "",
            "表格文件 (*.csv *.xls *.xlsx);;所有文件 (*.*)"
        )
        
        if file_paths:
            for file_path in file_paths:
                row = self.table_list.rowCount()
                self.table_list.insertRow(row)
                
                # 文件路径
                self.table_list.setItem(row, 0, QTableWidgetItem(file_path))
                
                # 标题（使用文件名，不带扩展名）
                title = os.path.splitext(os.path.basename(file_path))[0]
                self.table_list.setItem(row, 1, QTableWidgetItem(title))
                
                # 表头行数
                header_spin = create_spin_box(0, 2, 1)
                self.table_list.setCellWidget(row, 2, header_spin)
                
                # 操作按钮
                operations = QWidget()
                layout = QHBoxLayout(operations)
                layout.setContentsMargins(0, 0, 0, 0)
                
                preview_btn = QPushButton("预览")
                preview_btn.clicked.connect(lambda checked, r=row: self.preview_single_table(r))
                layout.addWidget(preview_btn)
                
                delete_btn = QPushButton("删除")
                delete_btn.clicked.connect(lambda checked, r=row: self.delete_table(r))
                layout.addWidget(delete_btn)
                
                self.table_list.setCellWidget(row, 3, operations)

    def paste_table(self):
        """粘贴表格"""
        dialog = PasteTableDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            try:
                df = dialog.get_dataframe()
                if df is not None and not df.empty:
                    row = self.table_list.rowCount()
                    self.table_list.insertRow(row)
                    
                    # 数据来源
                    self.table_list.setItem(row, 0, QTableWidgetItem("粘贴的数据"))
                    
                    # 标题
                    self.table_list.setItem(row, 1, QTableWidgetItem(f"表格 {row + 1}"))
                    
                    # 表头行数
                    header_spin = create_spin_box(0, 2, dialog.header_rows)
                    self.table_list.setCellWidget(row, 2, header_spin)
                    
                    # 操作按钮
                    operations = QWidget()
                    layout = QHBoxLayout(operations)
                    layout.setContentsMargins(0, 0, 0, 0)
                    
                    preview_btn = QPushButton("预览")
                    preview_btn.clicked.connect(lambda checked, r=row: self.preview_single_table(r))
                    layout.addWidget(preview_btn)
                    
                    delete_btn = QPushButton("删除")
                    delete_btn.clicked.connect(lambda checked, r=row: self.delete_table(r))
                    layout.addWidget(delete_btn)
                    
                    self.table_list.setCellWidget(row, 3, operations)
                    
                    # 保存数据
                    self.tables.append(df)
                    
                    QMessageBox.information(self, "成功", "表格已添加到列表")
            except Exception as e:
                QMessageBox.critical(self, "错误", str(e))

    def clear_tables(self):
        """清空表格列表"""
        self.table_list.setRowCount(0)
        self.tables.clear()

    def delete_table(self, row):
        """删除指定行的表格"""
        self.table_list.removeRow(row)
        if row < len(self.tables):
            self.tables.pop(row)

    def preview_single_table(self, row):
        """预览单个表格"""
        try:
            source = self.table_list.item(row, 0).text()
            caption = self.table_list.item(row, 1).text()
            header_rows = self.table_list.cellWidget(row, 2).value()
            
            if source == "粘贴的数据":
                df = self.tables[row].copy()
            else:
                df = load_dataframe(source, header_rows=header_rows)
            
            # 创建临时文件并预览
            temp_path = create_temp_docx()
            if not temp_path:
                return
            
            write_table_to_docx(
                df=df,
                output_path=temp_path,
                table_caption=caption,
                chinese_font=self.chinese_font_combo.currentText(),
                english_font=self.english_font_combo.currentText(),
                font_size=self.font_size_spin.value(),
                header_rows=header_rows,
                border_width=self.border_width_spin.value(),
                bold_caption=self.bold_caption_check.isChecked()
            )
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"预览失败: {str(e)}")

    def preview_tables(self):
        """预览所有表格"""
        if self.table_list.rowCount() == 0:
            QMessageBox.warning(self, "警告", "请先添加表格！")
            return
        
        try:
            # 收集所有表格数据
            tables = []
            captions = []
            header_rows_list = []
            
            for row in range(self.table_list.rowCount()):
                source = self.table_list.item(row, 0).text()
                caption = self.table_list.item(row, 1).text()
                header_rows = self.table_list.cellWidget(row, 2).value()
                
                if source == "粘贴的数据":
                    df = self.tables[row].copy()
                else:
                    df = load_dataframe(source, header_rows=header_rows)
                
                if df is not None and not df.empty:
                    tables.append(df)
                    captions.append(caption)
                    header_rows_list.append(header_rows)
            
            if not tables:
                QMessageBox.warning(self, "警告", "没有有效的表格数据！")
                return
            
            # 创建临时文件并预览
            temp_path = create_temp_docx()
            if not temp_path:
                return
            
            write_tables_to_docx(
                tables=tables,
                output_path=temp_path,
                captions=captions,
                font_name=self.chinese_font_combo.currentText(),
                font_size=self.font_size_spin.value(),
                header_rows=header_rows_list,
                separate_files=self.output_mode.currentText() == "每个表格单独保存"
            )
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"预览失败: {str(e)}")

    def export_tables(self):
        """导出所有表格"""
        if self.table_list.rowCount() == 0:
            QMessageBox.warning(self, "警告", "请先添加表格！")
            return
        
        try:
            # 获取保存路径
            if self.output_mode.currentText() == "每个表格单独保存":
                output_dir = QFileDialog.getExistingDirectory(self, "选择保存目录")
                if not output_dir:
                    return
            else:
                output_path, _ = QFileDialog.getSaveFileName(
                    self,
                    "保存文件",
                    "",
                    "Word 文档 (*.docx)"
                )
                if not output_path:
                    return
                if not output_path.endswith('.docx'):
                    output_path += '.docx'
            
            # 收集所有表格数据
            tables = []
            captions = []
            header_rows_list = []
            
            for row in range(self.table_list.rowCount()):
                source = self.table_list.item(row, 0).text()
                caption = self.table_list.item(row, 1).text()
                header_rows = self.table_list.cellWidget(row, 2).value()
                
                if source == "粘贴的数据":
                    df = self.tables[row].copy()
                else:
                    df = load_dataframe(source, header_rows=header_rows)
                
                if df is not None and not df.empty:
                    tables.append(df)
                    captions.append(caption)
                    header_rows_list.append(header_rows)
            
            if not tables:
                QMessageBox.warning(self, "警告", "没有有效的表格数据！")
                return
            
            # 导出文件
            if self.output_mode.currentText() == "每个表格单独保存":
                for i, (table, caption, header_rows) in enumerate(zip(tables, captions, header_rows_list)):
                    output_path = os.path.join(output_dir, f"{caption}.docx")
                    write_table_to_docx(
                        df=table,
                        output_path=output_path,
                        table_caption=caption,
                        chinese_font=self.chinese_font_combo.currentText(),
                        english_font=self.english_font_combo.currentText(),
                        font_size=self.font_size_spin.value(),
                        header_rows=header_rows,
                        border_width=self.border_width_spin.value(),
                        bold_caption=self.bold_caption_check.isChecked()
                    )
                QMessageBox.information(self, "成功", f"已导出 {len(tables)} 个文件到 {output_dir}")
            else:
                write_tables_to_docx(
                    tables=tables,
                    output_path=output_path,
                    captions=captions,
                    font_name=self.chinese_font_combo.currentText(),
                    font_size=self.font_size_spin.value(),
                    header_rows=header_rows_list,
                    separate_files=False
                )
                QMessageBox.information(self, "成功", f"文件已保存到: {output_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(get_translation("window_title"))
        self.setMinimumSize(800, 600)
        
        # 创建中心部件和布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 创建选项卡
        tab_widget = QTabWidget()
        tab_widget.setStyleSheet("""
            QTabBar {
                font-size: 14pt;
                font-family: SimHei;
            }
            QTabBar::tab {
                padding: 12px 20px;
                margin-top: 10px;
            }
            QTabWidget::pane {
                margin-top: 10px;
            }
        """)
        
        # 单表转换选项卡
        single_tab = QWidget()
        single_layout = QVBoxLayout(single_tab)
        
        # ===== 1. 数据导入区域 =====
        import_group = QGroupBox("数据导入")
        import_layout = QVBoxLayout()
        
        # 文件选择区域
        file_group = QHBoxLayout()
        self.file_path_edit = QLineEdit()
        file_group.addWidget(self.file_path_edit)
        
        select_file_btn = create_button("选择文件", self.select_file)
        file_group.addWidget(select_file_btn)
        
        paste_data_btn = create_button("粘贴数据", self.show_paste_dialog)
        file_group.addWidget(paste_data_btn)
        
        import_layout.addLayout(file_group)
        import_group.setLayout(import_layout)
        single_layout.addWidget(import_group)
        
        # ===== 2. 格式设置区域 =====
        format_group = QGroupBox("格式设置")
        format_layout = QGridLayout()
        format_layout.setHorizontalSpacing(100)  # 减小水平间距
        format_layout.setVerticalSpacing(10)    # 减小垂直间距
        
        # 中文字体
        format_layout.addWidget(QLabel("中文字体"), 0, 0)
        self.chinese_font_combo = create_combo_box(["宋体", "黑体", "楷体", "微软雅黑"])
        format_layout.addWidget(self.chinese_font_combo, 0, 1)
        
        # 英文字体
        format_layout.addWidget(QLabel("英文字体"), 0, 2)
        self.english_font_combo = create_combo_box(["Times New Roman", "Arial", "Calibri"])
        format_layout.addWidget(self.english_font_combo, 0, 3)
        
        # 字号
        format_layout.addWidget(QLabel("字号"), 1, 0)
        self.font_size_spin = create_spin_box(8, 72, 12)
        format_layout.addWidget(self.font_size_spin, 1, 1)
        
        # 三线表线宽
        format_layout.addWidget(QLabel("三线表线宽"), 1, 2)
        self.border_width_spin = QDoubleSpinBox()
        self.border_width_spin.setRange(0.5, 3.0)
        self.border_width_spin.setSingleStep(0.25)
        self.border_width_spin.setValue(1.0)
        format_layout.addWidget(self.border_width_spin, 1, 3)
        
        format_group.setLayout(format_layout)
        single_layout.addWidget(format_group)
        
        # ===== 3. 输出设置区域 =====
        output_group = QGroupBox("输出设置")
        output_layout = QGridLayout()
        output_layout.setHorizontalSpacing(100)  # 减小水平间距
        output_layout.setVerticalSpacing(10)    # 减小垂直间距
        
        # 表格标题
        output_layout.addWidget(QLabel("表格标题"), 0, 0)
        self.title_edit = QLineEdit()
        output_layout.addWidget(self.title_edit, 0, 1, 1, 2)
        
        # 标题加粗选项
        self.bold_caption_check = QCheckBox("标题加粗")
        self.bold_caption_check.setChecked(True)
        output_layout.addWidget(self.bold_caption_check, 0, 3)
        
        # 表头行数
        output_layout.addWidget(QLabel("表头行数"), 1, 0)
        self.header_rows_spin = create_spin_box(0, 2, 1)
        output_layout.addWidget(self.header_rows_spin, 1, 1)
        
        output_group.setLayout(output_layout)
        single_layout.addWidget(output_group)
        
        # ===== 4. 按钮区域 =====
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        preview_btn = create_button("预览", self.preview_table)
        button_layout.addWidget(preview_btn)
        
        export_btn = create_button("导出", self.export_table)
        button_layout.addWidget(export_btn)
        
        single_layout.addLayout(button_layout)
        
        # 批量转换选项卡
        batch_tab = BatchProcessDialog()
        
        # 添加选项卡到选项卡组件
        tab_widget.addTab(single_tab, "单表转换")
        tab_widget.addTab(batch_tab, "批量转换")
        
        main_layout.addWidget(tab_widget)
        
        # 设置字体
        app = QApplication.instance()
        font = QFont("SimHei", 12)
        app.setFont(font)

    def select_file(self):
        """选择输入文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            get_translation("select_file"),
            "",
            "Table Files (*.csv *.xls *.xlsx);;All Files (*.*)"
        )
        
        if file_path:
            try:
                df = load_dataframe(
                    file_path,
                    header_rows=self.header_rows_spin.value()
                )
                self.df = df
                self.file_path_edit.setText(file_path)
                show_message("info", get_translation("import_success"))
            except Exception as e:
                show_message("error", str(e))
    
    def show_paste_dialog(self):
        """显示粘贴数据对话框"""
        dialog = PasteTableDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            try:
                self.df = dialog.get_dataframe()
                self.header_rows_spin.setValue(dialog.header_rows)
                self.file_path_edit.clear()
                show_message("info", get_translation("paste_success"))
            except Exception as e:
                show_message("error", str(e))
    
    def preview_table(self):
        """预览表格"""
        if self.df is None or self.df.empty:
            show_message("warning", "请先加载或粘贴有效的表格数据！")
            return
        
        temp_path = None
        try:
            temp_path = create_temp_docx()
            if not temp_path:
                return
            
            # 重新读取数据以确保正确的表头行数
            if self.file_path_edit.text():
                preview_df = load_dataframe(
                    self.file_path_edit.text(),
                    header_rows=self.header_rows_spin.value()
                )
            else:
                preview_df = self.df.copy()
            
            write_table_to_docx(
                df=preview_df,
                output_path=temp_path,
                table_caption=self.title_edit.text(),
                chinese_font=self.chinese_font_combo.currentText(),
                english_font=self.english_font_combo.currentText(),
                font_size=self.font_size_spin.value(),
                header_rows=self.header_rows_spin.value(),
                border_width=self.border_width_spin.value(),
                bold_caption=self.bold_caption_check.isChecked()
            )
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"预览失败: {str(e)}")
            if temp_path:
                cleanup_temp_file(temp_path)
    
    def export_table(self):
        """导出表格"""
        if self.df is None or self.df.empty:
            show_message("warning", "请先加载或粘贴有效的表格数据！")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存DOCX文件",
            "",
            "Word Documents (*.docx);;All Files (*.*)"
        )
        
        if file_path:
            try:
                export_df = self.df.copy()
                
                write_table_to_docx(
                    df=export_df,
                    output_path=file_path,
                    table_caption=self.title_edit.text(),
                    chinese_font=self.chinese_font_combo.currentText(),
                    english_font=self.english_font_combo.currentText(),
                    font_size=self.font_size_spin.value(),
                    header_rows=self.header_rows_spin.value(),
                    border_width=self.border_width_spin.value(),
                    bold_caption=self.bold_caption_check.isChecked()
                )
                show_message("info", "导出成功")
            except Exception as e:
                show_message("error", str(e))


def main():
    """主函数"""
    app = QApplication(sys.argv)
    
    # 设置应用程序信息
    app.setApplicationName("Research Table Converter")
    app.setApplicationVersion("1.0.0")
    app.setOrganizationName("JiehaoZhang")
    app.setOrganizationDomain("github.com/JiehaoZhang1993")
    
    # 创建主窗口
    window = MainWindow()
    window.show()
    
    # 运行应用程序
    sys.exit(app.exec_()) 

if __name__ == '__main__':
    main() 