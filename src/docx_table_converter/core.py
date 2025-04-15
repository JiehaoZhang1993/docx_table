"""
Core functionality for converting tables to DOCX format.
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import io
import docx.shared
import docx.enum.text
import docx.enum.style
import docx.enum.table


def read_table_from_file(file_path, sheet_name=None, header_rows=1):
    """
    Read a table from a file (CSV, XLS, XLSX) and return it as a pandas DataFrame.
    
    Parameters:
    -----------
    file_path : str
        Path to the input file
    sheet_name : str or int, optional
        Name or index of the sheet to read (for Excel files). If None, reads the first sheet.
    header_rows : int, optional
        Number of header rows in the table (default: 1)
    
    Returns:
    --------
    pandas.DataFrame
        The table data as a DataFrame
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.csv':
        return pd.read_csv(file_path, header=list(range(header_rows)))
    elif file_ext in ['.xls', '.xlsx']:
        # If sheet_name is None, read the first sheet
        if sheet_name is None:
            return pd.read_excel(file_path, header=list(range(header_rows)))
        return pd.read_excel(file_path, sheet_name=sheet_name, header=list(range(header_rows)))
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def parse_clipboard_data(text, header_rows=1):
    """
    Parse table data from clipboard text.
    
    Parameters:
    -----------
    text : str
        Table data as text (tab or comma separated)
    header_rows : int, optional
        Number of header rows (default: 1)
    
    Returns:
    --------
    pandas.DataFrame
        The table data as a DataFrame
    """
    # Try to detect the delimiter
    sample_line = text.split('\n')[0]
    if '\t' in sample_line:
        delimiter = '\t'
    else:
        delimiter = ','
    
    # Convert text to DataFrame
    buffer = io.StringIO(text)
    df = pd.read_csv(buffer, sep=delimiter, header=list(range(header_rows-1)))
    
    return df


def write_table_to_docx(df, output_path, table_caption="Table 1", headers=None,
                        chinese_font='宋体', english_font='Times New Roman', font_size=12,
                        header_rows=1, mode='append', border_width=1, bold_caption=True):
    """
    Write a pandas DataFrame as a formatted table to a Word document.
    If the document exists, it will append the table to it or overwrite it based on mode parameter.
    
    Parameters:
    -----------
    df : pandas.DataFrame
        The DataFrame to be written as a table
    output_path : str
        Path where the Word document should be saved
    table_caption : str, optional
        Caption for the table (default: "Table 1")
    headers : list, optional
        Custom headers for the table columns (default: None)
    chinese_font : str, optional
        Font to use for Chinese characters (default: '宋体')
    english_font : str, optional
        Font to use for English characters (default: 'Times New Roman')
    font_size : int, optional
        Font size in points (default: 12)
    header_rows : int, optional
        Number of header rows in the table (default: 1)
    mode : str, optional
        Mode to use when file exists: 'append' (add to existing document) or 'overwrite' (create new document)
        (default: 'append')
    border_width : float, optional
        Width of table borders in points (default: 1)
    bold_caption : bool, optional
        Whether to make the table caption bold (default: True)
    """
    # Check if document exists
    if os.path.exists(output_path) and mode == 'append':
        doc = Document(output_path)
    else:
        # Create a new document
        doc = Document()
        
        # Set page margins (2.5 cm)
        sections = doc.sections
        for section in sections:
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = english_font
        font.size = Pt(font_size)
        # Set east asian font (Chinese)
        if hasattr(style, '_element'):
            rPr = style._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), chinese_font)
    
    # Add table caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_run = caption.add_run(table_caption)
    caption_run.bold = bold_caption
    caption_run.font.name = english_font
    caption_run.font.size = Pt(font_size)
    # Set east asian font for caption
    if hasattr(caption_run, '_element'):
        rPr = caption_run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), chinese_font)
    
    # Check if DataFrame has MultiIndex columns
    has_multi_level_columns = isinstance(df.columns, pd.MultiIndex)
    
    # Initialize num_header_rows to avoid UnboundLocalError
    num_header_rows = header_rows
    
    # Prepare data for table
    df_for_table = df.copy()
    
    # Reset index if it's not a default RangeIndex
    if not isinstance(df_for_table.index, pd.RangeIndex):
        df_for_table = df_for_table.reset_index(drop=True)
    
    table_cols = len(df_for_table.columns)
    if headers is None:
        headers = df_for_table.columns
    
    # Handle multi-level columns
    if has_multi_level_columns:
        num_header_rows = len(df_for_table.columns.levels)
        rows = min(len(df_for_table) + num_header_rows, 1000)  # Limit rows to prevent index errors
        table = doc.add_table(rows=rows, cols=table_cols)
        
        # Add multi-level headers
        for level in range(num_header_rows):
            # Track cells to merge for each level
            merge_ranges = []
            start_idx = 0
            current_val = None
            
            # Identify ranges to merge
            for col_idx, col_val in enumerate(df_for_table.columns.get_level_values(level)):
                if col_val != current_val:
                    if col_idx > start_idx and current_val is not None:
                        merge_ranges.append((start_idx, col_idx - 1))
                    current_val = col_val
                    start_idx = col_idx
                    
                    # Only add text to the first cell of each merged range
                    cell = table.cell(level, col_idx)
                    cell.text = str(col_val)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.name = english_font
                    run.font.size = Pt(font_size)
                    # Set east asian font for cell text
                    if hasattr(run, '_element'):
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), chinese_font)
            
            # Add the last range if needed
            if start_idx < table_cols - 1 and current_val is not None:
                merge_ranges.append((start_idx, table_cols - 1))
            
            # Perform the merges
            for start, end in merge_ranges:
                if start != end:  # Only merge if there's more than one cell
                    table.cell(level, start).merge(table.cell(level, end))
        
        # Adjust data row start index
        data_row_start = num_header_rows

    else:
        rows = min(len(df_for_table) + 1, 1000)  # Limit rows to prevent index errors
        table = doc.add_table(rows=rows, cols=table_cols)
        
        # Add single-level headers
        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = str(header)
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.name = english_font
            run.font.size = Pt(font_size)
            # Set east asian font for cell text
            if hasattr(run, '_element'):
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), chinese_font)
        
        # Adjust data row start index
        data_row_start = 1
    
    table.style = 'Table Grid'
    
    # Handle table style safely
    if hasattr(table, 'style') and hasattr(table.style, 'element'):
        borders = table.style.element.xpath('//w:tblBorders')
        if borders and len(borders) > 0:
            parent = borders[0].getparent()
            if parent is not None:
                parent.remove(borders[0])
    
    # Add data
    for i, (idx, row) in enumerate(df_for_table.iterrows()):
        if i >= rows - data_row_start:  # Check if we're exceeding table dimensions
            break
        for j, value in enumerate(row):
            cell = table.cell(i + data_row_start, j)
            
            # Check if the value contains formatting like subscripts or superscripts
            value_str = str(value)
            
            # Clear existing text
            cell.text = ""
            paragraph = cell.paragraphs[0]
            
            # Handle subscripts (e.g., $_{text}$)
            sub_pattern = re.compile(r'\$_{(.+?)}\$')
            parts = sub_pattern.split(value_str)
            
            for k, part in enumerate(parts):
                if k % 2 == 0:  # Regular text
                    if part:
                        run = paragraph.add_run(part)
                        run.font.name = english_font
                        run.font.size = Pt(font_size)
                        # Set east asian font for cell text
                        if hasattr(run, '_element'):
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), chinese_font)
                else:  # Subscript text
                    run = paragraph.add_run(part)
                    run.font.name = english_font
                    run.font.size = Pt(font_size)
                    run.font.subscript = True
                    # Set east asian font for subscript text
                    if hasattr(run, '_element'):
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), chinese_font)
            
            # If no formatting was found, add the text normally
            if len(parts) == 1:
                cell.text = value_str
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = english_font
                    run.font.size = Pt(font_size)
                    # Set east asian font for normal cell text
                    if hasattr(run, '_element'):
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), chinese_font)
    
    # Apply three-line table style with no vertical lines
    border_size = int(border_width * 8)  # Convert points to eighths of a point
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Determine border style based on row position and header rows
            top_border = 'none'
            bottom_border = 'none'
            
            if has_multi_level_columns:
                if row_idx == 0:  # First header row
                    top_border = 'single'
                elif row_idx == num_header_rows - 1:  # Last header row
                    bottom_border = 'single'
                elif row_idx == len(table.rows) - 1:  # Last data row
                    bottom_border = 'single'
            else:
                if row_idx == 0:  # Header row
                    top_border = 'single'
                    bottom_border = 'single'
                elif row_idx == len(table.rows) - 1:  # Last data row
                    bottom_border = 'single'
            
            borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                <w:top w:val="{top_border}" w:sz="{border_size}"/>
                <w:bottom w:val="{bottom_border}" w:sz="{border_size}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tcBorders>'''
            existing_borders = tcPr.xpath('./w:tcBorders')
            if existing_borders:
                tcPr.remove(existing_borders[0])
            tcBorders = parse_xml(borders_xml)
            tcPr.append(tcBorders)
    
    # Add page break after the table
    doc.add_paragraph().add_run().add_break()
    
    # Save document
    doc.save(output_path)
    # Open saved file with word or default application
    if os.name == 'nt':
        os.startfile(os.path.abspath(output_path))
    elif os.name == 'posix':
        os.system(f'xdg-open "{os.path.abspath(output_path)}"')
    else:
        print(f"File saved at {output_path}. Please open it manually.")
    return None


def write_tables_to_docx(tables, captions, output_path, descriptions=None, font_name="Times New Roman",
                     font_size=10.5, header_rows=None, index_cols=None, separate_files=False):
    """
    将多个表格写入DOCX文件。

    Args:
        tables: list，包含多个pandas DataFrame或文件路径
        captions: list，表格标题列表
        output_path: str，输出文件路径
        descriptions: list，可选，表格描述列表，默认为None
        font_name: str，可选，英文字体名称，默认为Times New Roman
        font_size: float，可选，字体大小，默认为10.5
        header_rows: list，可选，每个表格的表头行数列表，默认为None
        index_cols: list，可选，每个表格的索引列名列表，默认为None
        separate_files: bool，可选，是否将每个表格保存为单独的文件，默认为False

    Returns:
        list: 生成的文件路径列表
    """
    if not isinstance(tables, list):
        raise ValueError("tables参数必须是列表类型")
    if not isinstance(captions, list):
        raise ValueError("captions参数必须是列表类型")
    if len(tables) != len(captions):
        raise ValueError("tables和captions的长度必须相同")
    
    # 初始化可选参数
    if descriptions is None:
        descriptions = [""] * len(tables)
    if header_rows is None:
        header_rows = [1] * len(tables)
    if index_cols is None:
        index_cols = [None] * len(tables)
    
    # 验证可选参数的长度
    if len(descriptions) != len(tables):
        raise ValueError("descriptions的长度必须与tables相同")
    if len(header_rows) != len(tables):
        raise ValueError("header_rows的长度必须与tables相同")
    if len(index_cols) != len(tables):
        raise ValueError("index_cols的长度必须与tables相同")
    
    generated_files = []
    
    if separate_files:
        # 为每个表格创建单独的文件
        
        for i, (table, caption) in enumerate(zip(tables, captions)):
            file_path = f"{output_path}_{i + 1}.docx"
            
            # 读取数据（如果是文件路径）
            if isinstance(table, str):
                df = read_table_from_file(table, header_rows=header_rows[i])
            else:
                df = table.copy()
            
            # 写入单个表格
            write_table_to_docx(
                df=df,
                output_path=file_path,
                table_caption=caption,
                chinese_font='宋体',
                english_font=font_name,
                font_size=font_size,
                mode='overwrite'
            )
            generated_files.append(file_path)
    else:
        path = os.path.abspath(output_path)
        for i, (table, caption) in enumerate(zip(tables, captions)):
            # 读取数据（如果是文件路径）
            if isinstance(table, str):
                df = read_table_from_file(table, header_rows=header_rows[i])
            else:
                df = table.copy()
            write_table_to_docx(df=df,
                output_path=path,
                table_caption=caption,
                chinese_font='宋体',
                english_font=font_name,
                font_size=font_size,
                mode='append'
            )
   
        generated_files.append(output_path)
    
    return generated_files 