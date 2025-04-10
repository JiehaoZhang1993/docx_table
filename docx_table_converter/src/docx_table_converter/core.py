"""
Core functionality for converting tables to DOCX format.
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
import re
import io
import openpyxl
import xlrd


def read_table_from_file(file_path, sheet_name=None, header_rows=1):
    """
    Read a table from a file (CSV, XLS, XLSX) and return it as a pandas DataFrame.
    
    Parameters:
    -----------
    file_path : str
        Path to the input file
    sheet_name : str or int, optional
        Name or index of the sheet to read (for Excel files). If None, reads the first sheet.
    header_rows : int, optional
        Number of header rows in the table (default: 1)
    
    Returns:
    --------
    pandas.DataFrame
        The table data as a DataFrame
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.csv':
        return pd.read_csv(file_path, header=list(range(header_rows)))
    elif file_ext in ['.xls', '.xlsx']:
        # If sheet_name is None, read the first sheet
        if sheet_name is None:
            return pd.read_excel(file_path, header=list(range(header_rows)))
        return pd.read_excel(file_path, sheet_name=sheet_name, header=list(range(header_rows)))
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def parse_clipboard_data(text, header_rows=1):
    """
    Parse table data from clipboard text.
    
    Parameters:
    -----------
    text : str
        Table data as text (tab or comma separated)
    header_rows : int, optional
        Number of header rows (default: 1)
    
    Returns:
    --------
    pandas.DataFrame
        The table data as a DataFrame
    """
    # Try to detect the delimiter
    sample_line = text.split('\n')[0]
    if '\t' in sample_line:
        delimiter = '\t'
    else:
        delimiter = ','
    
    # Convert text to DataFrame
    buffer = io.StringIO(text)
    df = pd.read_csv(buffer, sep=delimiter, header=list(range(header_rows-1)))
    
    return df


def write_table_to_docx(df, output_path, table_caption="Table", table_description="",
                       headers=None, font_name='Times New Roman', font_size=12,
                       include_index=False, mode='append', caption_bold=True,
                       header_bold=True, table_width=None, column_widths=None,
                       cell_padding=0.1, border_width=1.0, border_color='000000',
                       header_background_color=None, caption_alignment='left',
                       table_alignment='center', page_margins=(2.5, 2.5, 2.5, 2.5),
                       page_size='A4', page_orientation='portrait',
                       page_break_after=True, special_formatting=True):
    """
    Write a pandas DataFrame as a formatted table to a Word document.
    
    Parameters:
    -----------
    df : pandas.DataFrame
        The DataFrame to be written as a table
    output_path : str
        Path where the Word document should be saved
    table_caption : str, optional
        Caption number/identifier for the table (default: "Table")
    table_description : str, optional
        Description text for the table (default: "")
    headers : list, optional
        Custom headers for the table. If None, DataFrame column names will be used
    font_name : str, optional
        Font to use in the document (default: 'Times New Roman')
    font_size : int, optional
        Font size in points (default: 12)
    include_index : bool, optional
        Whether to include the DataFrame index in the table (default: False)
    mode : str, optional
        Mode to use when file exists: 'append' or 'overwrite' (default: 'append')
    caption_bold : bool, optional
        Whether to make the caption bold (default: True)
    header_bold : bool, optional
        Whether to make the header bold (default: True)
    table_width : float, optional
        Table width in inches (default: None, auto-width)
    column_widths : list, optional
        List of column widths in inches (default: None, equal width)
    cell_padding : float, optional
        Cell padding in inches (default: 0.1)
    border_width : float, optional
        Border width in points (default: 1.0)
    border_color : str, optional
        Border color in hex format (default: '000000')
    header_background_color : str, optional
        Header background color in hex format (default: None)
    caption_alignment : str, optional
        Caption alignment: 'left', 'center', or 'right' (default: 'left')
    table_alignment : str, optional
        Table alignment: 'left', 'center', or 'right' (default: 'center')
    page_margins : tuple, optional
        Page margins in cm (left, right, top, bottom) (default: (2.5, 2.5, 2.5, 2.5))
    page_size : str, optional
        Page size: 'A4', 'Letter', etc. (default: 'A4')
    page_orientation : str, optional
        Page orientation: 'portrait' or 'landscape' (default: 'portrait')
    page_break_after : bool, optional
        Whether to add a page break after the table (default: True)
    special_formatting : bool, optional
        Whether to process special formatting in text (default: True)
    """
    # Check if document exists
    if os.path.exists(output_path) and mode == 'append':
        doc = Document(output_path)
    else:
        doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(page_margins[0])
        section.right_margin = Cm(page_margins[1])
        section.top_margin = Cm(page_margins[2])
        section.bottom_margin = Cm(page_margins[3])
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    if hasattr(style, '_element') and hasattr(style._element, 'rPr'):
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    
    # Add table caption
    caption = doc.add_paragraph()
    caption_align = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT
    }
    caption.alignment = caption_align.get(caption_alignment, WD_ALIGN_PARAGRAPH.LEFT)
    
    caption_text = caption.add_run(f'{table_caption}. ')
    caption_text.bold = caption_bold
    caption_text.font.name = font_name
    caption_text.font.size = Pt(font_size)
    
    if table_description:
        desc_text = caption.add_run(table_description)
        desc_text.bold = False
        desc_text.font.name = font_name
        desc_text.font.size = Pt(font_size)
    
    # Prepare data for table
    if include_index:
        df_with_index = df.copy()
        df_with_index = df_with_index.reset_index()
        table_cols = len(df_with_index.columns)
        if headers is None:
            headers = df_with_index.columns
    else:
        table_cols = len(df.columns)
        if headers is None:
            headers = df.columns
    
    data_df = df_with_index if include_index else df
    
    # Create table
    rows = min(len(data_df) + 1, 1000)  # Limit rows to prevent index errors
    table = doc.add_table(rows=rows, cols=table_cols)
    table.style = 'Table Grid'
    
    # Set table alignment
    if table_alignment == 'center':
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif table_alignment == 'right':
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add headers
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = str(header)
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.bold = header_bold
        run.font.name = font_name
        run.font.size = Pt(font_size)
        
        if header_background_color:
            cell._tc.get_or_add_tcPr().append(parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{header_background_color}"/>'
            ))
    
    # Add data
    for i, (idx, row) in enumerate(data_df.iterrows()):
        if i >= rows - 1:  # Check if we're exceeding table dimensions
            break
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            
            if special_formatting:
                # Handle special formatting
                value_str = str(value)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                
                # Handle subscripts (e.g., $_{text}$)
                sub_pattern = re.compile(r'\$_{(.+?)}\$')
                parts = sub_pattern.split(value_str)
                
                for k, part in enumerate(parts):
                    if k % 2 == 0:  # Regular text
                        if part:
                            run = paragraph.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                    else:  # Subscript text
                        run = paragraph.add_run(part)
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.subscript = True
            else:
                cell.text = str(value)
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
    
    # Apply three-line table style with no vertical frame lines
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Determine border style based on row position
            top_border = 'none'
            bottom_border = 'none'
            
            if row_idx == 0:  # Header row
                top_border = 'single'
                bottom_border = 'single'
            elif row_idx == len(table.rows) - 1:  # Last data row
                bottom_border = 'single'
            
            # Apply borders - with no vertical frame lines
            borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                <w:top w:val="{top_border}" w:sz="{int(border_width * 8)}" w:color="{border_color}"/>
                <w:bottom w:val="{bottom_border}" w:sz="{int(border_width * 8)}" w:color="{border_color}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tcBorders>'''
            
            existing_borders = tcPr.xpath('./w:tcBorders')
            if existing_borders:
                tcPr.remove(existing_borders[0])
            tcBorders = parse_xml(borders_xml)
            tcPr.append(tcBorders)
    
    # Set table width and column widths
    if table_width or column_widths:
        table.allow_autofit = False
        if table_width:
            table.width = Inches(table_width)
        if column_widths:
            for i, width in enumerate(column_widths):
                if i < table_cols:
                    table.columns[i].width = Inches(width)
    
    # Add page break if requested
    if page_break_after:
        doc.add_paragraph().add_run().add_break()
    
    # Save document
    doc.save(output_path)
    
    return None 